from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path(
    r"C:\Users\creat\Documents\app meeting recall\docs"
    r"\The AI App Building Playbook - Revised App Store Edition.docx"
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.08
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size in [("Heading 1", 22), ("Heading 2", 15), ("Heading 3", 12)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(12)
        styles[name].paragraph_format.space_after = Pt(6)

    styles["Subtitle"].font.name = "Aptos"
    styles["Subtitle"].font.size = Pt(12)
    styles["Subtitle"].font.italic = True


def add_page_break(doc):
    doc.add_page_break()


def add_h1(doc, text, subtitle=None):
    add_page_break(doc)
    p = doc.add_paragraph(style="Heading 1")
    p.add_run(text)
    if subtitle:
        doc.add_paragraph(subtitle, style="Subtitle")


def add_h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def add_h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def add_p(doc, text=""):
    return doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_prompt(doc, title, text):
    add_h3(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F6FC")
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    doc.add_paragraph()


def add_case(doc, title, text):
    add_h3(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7F7")
    cell.paragraphs[0].add_run(text)
    doc.add_paragraph()


def add_checklist(doc, title, items):
    add_h3(doc, title)
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run("[ ] ").bold = True
        p.add_run(item)


def add_front_matter(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("The AI App Building Playbook")
    r.bold = True
    r.font.size = Pt(30)
    r.font.name = "Aptos Display"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "A practical guide for turning an app idea into a real product with "
        "ChatGPT, Codex, and AI-assisted development"
    )
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Revised App Store Edition").italic = True

    add_page_break(doc)
    add_h2(doc, "About This Edition")
    add_p(
        doc,
        "This edition expands the original playbook from prototype guidance into a "
        "complete path from idea to real app-store submission. It is written for "
        "founders, designers, creators, students, and non-developers who want to "
        "use ChatGPT and Codex as serious product-building partners.",
    )
    add_p(
        doc,
        "The goal is not to make you a senior engineer overnight. The goal is to "
        "help you think clearly, prompt well, reduce product risk, test honestly, "
        "and bring a focused app idea to life without wasting months in the wrong "
        "direction.",
    )
    add_p(
        doc,
        "The examples in this edition draw from the real process of building "
        "Meeting Recall, a local-first meeting recorder designed to save audio "
        "files and help users open NotebookLM for summaries and insights. That "
        "project surfaced lessons about product scope, native mobile limitations, "
        "file storage, permissions, beta testing, App Review, Google Play "
        "production access, and the difference between building a cool prototype "
        "and shipping a trustworthy product.",
    )

    add_h2(doc, "Recommended Amazon/KDP Format")
    add_p(
        doc,
        "If you sell this as an ebook or paperback later, keep the format simple "
        "and professional. Amazon Kindle works best with clean document structure, "
        "real headings, normal paragraphs, and minimal complex layout.",
    )
    add_bullets(
        doc,
        [
            "Use a 6 x 9 inch trim size for paperback. It is common, readable, and cost-effective.",
            "Use real Word heading styles so a Kindle table of contents can be generated cleanly.",
            "Avoid text boxes, floating objects, complicated multi-column layouts, and decorative sidebars.",
            "Use page breaks before each major chapter.",
            "Use simple prompt boxes with tables or shaded paragraphs, not images of text.",
            "Export a print PDF for paperback and a separate reflowable EPUB or DOCX upload for Kindle.",
            "Keep screenshots optional. If included, use high-resolution images and explain the workflow in text.",
            "Add front matter: title page, copyright, disclaimer, who this is for, and how to use the book.",
            "Add back matter: resources, prompt library, launch checklist, and next steps.",
        ],
    )
    add_p(
        doc,
        'Suggested sale-page subtitle: "From Idea to App Store with ChatGPT, Codex, and AI-Assisted Development."',
    )


def add_main_content(doc):
    add_h1(doc, "Part 1: Think Before You Build", "The product decisions that make AI useful instead of chaotic.")
    add_h2(doc, "The Biggest Mistake Beginners Make")
    add_p(
        doc,
        "Most people open ChatGPT or Codex and ask it to build an app too early. "
        "That usually creates a messy prototype because the AI has no real product "
        "boundaries. It does not know who the app is for, what matters most, what "
        "should be avoided, or what technical risks could sink the project.",
    )
    add_p(
        doc,
        "The better approach is to use AI as a product partner before you use it "
        "as a coding tool. Make it help you think. Make it ask questions. Make it "
        "challenge the scope. Make it identify risky assumptions.",
    )
    add_h2(doc, "The App Clarity Stack")
    add_bullets(
        doc,
        [
            "Problem: What pain does this app reduce?",
            "Audience: Who feels that pain often enough to care?",
            "Moment: When does the user reach for this app?",
            "Outcome: What does the user get after using it?",
            "Trust: What must never go wrong?",
            "MVP: What is the smallest version that still delivers the core outcome?",
            "Not Included: What are you intentionally refusing to build?",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Clarify The Idea",
        "I have an app idea: [describe it]. Help me clarify the problem, target user, "
        "emotional reason to exist, core workflow, MVP, and what I should avoid building "
        "too early. Be direct and challenge weak assumptions.",
    )
    add_prompt(
        doc,
        "Prompt: Define What Must Never Break",
        "For this app idea, identify the 5 things that must never go wrong for users to "
        "trust it. For each one, explain the product risk, technical risk, and UX risk.",
    )
    add_case(
        doc,
        "Case Study: Meeting Recall Product Lock",
        "Meeting Recall became stronger by refusing to become an AI assistant, transcription "
        "tool, collaboration platform, or cloud storage product. Its role became clear: "
        "record, save, open NotebookLM, upload, and get insights. That boundary shaped every "
        "design and engineering decision.",
    )

    add_h1(doc, "Part 2: Turn The Idea Into A Product Brief", "Give Codex durable context before asking it to code.")
    add_h2(doc, "Why Documentation Matters")
    add_p(
        doc,
        "AI coding tools perform better when they have durable context. A prompt is "
        "temporary. A project document gives the AI something to keep returning to. "
        "If you care about scope, UX, technical decisions, or launch rules, write them "
        "down inside the repo.",
    )
    add_h2(doc, "Core Documents To Create")
    add_bullets(
        doc,
        [
            "PRODUCT_LOCK.md: what the product is and is not.",
            "UX_RULES.md: copy, flow, emotional tone, and friction rules.",
            "COMPONENT_SYSTEM.md: buttons, typography, spacing, states, and icons.",
            "TECHNICAL_EXPECTATIONS.md: platform behavior, storage, permissions, APIs, and reliability rules.",
            "IMPLEMENTATION_PLAN.md: phases and build order.",
            "QA_TEST_PLAN.md: real flows to test before launch.",
            "FEATURE_DONE_CHECKLIST.md: what done actually means.",
            "SCREEN_STATE_MATRIX.md: loading, empty, success, error, and interrupted states.",
            "ARCHITECTURE_DECISIONS.md: why you chose one technical path over another.",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Create A Product Lock",
        "Create PRODUCT_LOCK.md for this app. It should define the app, target user, core "
        "workflow, MVP scope, out-of-scope features, success definition, and product "
        "principles. Be strict about avoiding feature creep.",
    )
    add_prompt(
        doc,
        "Prompt: Create A QA Plan",
        "Create QA_TEST_PLAN.md for this app. Include first launch, permissions, core "
        "workflow, saved item flow, old item flow, share/export flow, error states, app "
        "restart, and real-device testing. Make it practical for beta testers.",
    )

    add_h1(doc, "Part 3: Prompt Codex Like A Product Lead", "Good prompts reduce chaos and protect the user experience.")
    add_h2(doc, "The Codex Prompt Formula")
    add_bullets(
        doc,
        [
            "Goal: what should be true when the task is done.",
            "Do not: what Codex must avoid changing or adding.",
            "Requirements: specific behavior, copy, screens, and edge cases.",
            "References: docs or files Codex should follow.",
            "Validation: tests, typecheck, build, or manual QA steps.",
            "Summary: what Codex should report back.",
        ],
    )
    add_prompt(
        doc,
        "Reusable Codex Task Prompt",
        "Goal:\n[Describe the outcome.]\n\nDo not:\n- Do not add unrelated features.\n"
        "- Do not refactor unrelated code.\n- Do not change working flows unless needed.\n\n"
        "Requirements:\n1. [Requirement]\n2. [Requirement]\n3. [Requirement]\n\n"
        "Follow these docs:\n- /docs/PRODUCT_LOCK.md\n- /docs/UX_RULES.md\n"
        "- /docs/TECHNICAL_EXPECTATIONS.md\n\nAfter completion:\n"
        "- run typecheck/build if available\n- summarize files changed\n- explain what to test",
    )
    add_h2(doc, "Rules That Save Projects")
    add_bullets(
        doc,
        [
            "Ask for one meaningful change at a time.",
            "Use “Do not implement...” to prevent accidental scope creep.",
            "Tell Codex what must be preserved.",
            "Ask it to update docs when behavior changes.",
            "Ask for exact testing steps after every major change.",
            "When something breaks, ask for root cause before adding features.",
            "When a flow feels wrong, describe the emotional problem, not just the UI bug.",
        ],
    )

    add_h1(doc, "Part 4: Build In Phases, Not Feature Piles", "The practical roadmap from shell to store.")
    add_h2(doc, "Recommended Build Order")
    add_numbered(
        doc,
        [
            "Project audit and stack validation.",
            "App shell and navigation.",
            "Theme and reusable components.",
            "Core screens with placeholders.",
            "Technical spikes for risky workflows.",
            "Production version of the core workflow.",
            "Saved item/history flow.",
            "Primary handoff/share/export flow.",
            "Settings, onboarding, and permissions.",
            "Error states and recovery states.",
            "Beta testing cleanup.",
            "Store assets, screenshots, and submission.",
        ],
    )
    add_checklist(
        doc,
        "Phase Gate Checklist",
        [
            "Can the user complete the core workflow end to end?",
            "Did we test on a real device?",
            "Did we document known limitations?",
            "Did we update QA steps?",
            "Are debug screens or spike controls still visible?",
            "Does the app still match the product lock?",
            "Is this change worth the added complexity?",
        ],
    )

    add_h1(doc, "Part 5: Technical Spikes Before Beautiful UI", "Validate what can sink the app before polishing it.")
    add_p(
        doc,
        "A spike is a small technical experiment that answers a risky question. It is "
        "not final UI. It is not the polished product. It is how you avoid spending "
        "weeks designing around an assumption that turns out to be false.",
    )
    add_bullets(
        doc,
        [
            "Can the app record audio reliably on real devices?",
            "Can recordings survive long sessions?",
            "Can files be saved somewhere users can find?",
            "Can files be shared with the correct MIME type?",
            "Can the target external app actually accept the file?",
            "Can permissions be requested without creating review problems?",
            "Does the platform support the flow you imagined?",
        ],
    )
    add_case(
        doc,
        "Meeting Recall Spike Lessons",
        "The file accessibility spike showed that Android could create visible files, but "
        "exported audio could become 0B if bytes were not written correctly. Rename experiments "
        "caused memory warnings. The final decision was to create the correct filename during "
        "initial save and defer risky post-save rename behavior.",
    )
    add_prompt(
        doc,
        "Prompt: Define Spikes",
        "Identify the highest-risk technical assumptions in this app. Create IMPLEMENTATION_SPIKES.md "
        "with each spike goal, test steps, pass/fail conditions, platform risks, and what decision "
        "the spike should unlock. Order the spikes by product risk.",
    )

    add_h1(doc, "Part 6: UX Is Not Decoration", "Build trust by removing uncertainty.")
    add_p(
        doc,
        "A user-friendly app does not just look clean. It reduces doubt. The user should "
        "not wonder what to do next, whether something saved, where a file went, or whether "
        "a button is safe to press.",
    )
    add_bullets(
        doc,
        [
            "Loading: what is happening right now?",
            "Empty: what can I do next?",
            "Success: did it work?",
            "Error: what went wrong in plain language?",
            "Interrupted: did I lose anything?",
            "Recovery: what action can I take now?",
        ],
    )
    add_case(
        doc,
        "File Readiness Gate",
        "Some Android testers could see audio files but could not select them in NotebookLM. "
        "Instead of adding scary helper copy, the product decision was to disable Open NotebookLM "
        "and Share until the saved file passed readiness checks. The ready state stayed confident. "
        "The failed state explained the problem only if needed.",
    )
    add_case(
        doc,
        "App Review Permission Lesson",
        "Meeting Recall was rejected because the custom microphone explainer used a button labeled "
        "“Allow Microphone Access.” The fix was to change the title to “Microphone access,” explain "
        "that the app uses the microphone to record meetings, and make the CTA “Continue.” The native "
        "OS prompt is where the user chooses allow or deny.",
    )
    add_prompt(
        doc,
        "Prompt: Add A Readiness Gate",
        "Add a readiness gate before enabling [critical action]. The screen should show “Preparing...” "
        "while validation runs, enable the primary action only after validation passes, and show a calm "
        "failed state with Try Again if validation fails. Do not add normal-state warning copy.",
    )

    add_h1(doc, "Part 7: Native Mobile Reality", "The difference between “it works for me” and “it works for users.”")
    add_p(
        doc,
        "Simulators are useful, but they do not fully represent file pickers, permissions, "
        "audio interruptions, storage providers, background behavior, or app review conditions. "
        "If the app depends on microphone, storage, camera, Bluetooth, notifications, calendar, "
        "or location, test on real devices early.",
    )
    add_h2(doc, "Android vs iOS Differences")
    add_bullets(
        doc,
        [
            "Android has many device brands, file providers, storage behaviors, and picker inconsistencies.",
            "iOS is more controlled, but sandboxing can make files less browsable unless sharing/export is designed well.",
            "Android may show a file but another app may gray it out due to MIME or provider behavior.",
            "iOS often prefers share/open-in flows over expecting users to browse app storage.",
            "Both platforms need clear permission copy and recovery states.",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Stack Validation",
        "Review the current codebase and validate whether the stack can support the product requirements. "
        "Be direct about limitations, native module needs, app store readiness, and recommended path forward.",
    )

    add_h1(doc, "Part 8: Local Storage, Sharing, And External Workflows", "If files are part of the product, files are part of the UX.")
    add_p(
        doc,
        "For apps that record, export, import, or upload files, storage is part of the user "
        "experience. Users need confidence that the file exists, has the right name, can be found, "
        "and can be used by the target app.",
    )
    add_bullets(
        doc,
        [
            "Use human-readable filenames.",
            "Create the correct filename during initial save whenever possible.",
            "Validate file size before marking a save successful.",
            "Validate playback or readability when feasible.",
            "Do not create fake success states.",
            "Make Share a reliable fallback for platform picker weirdness.",
            "Document platform differences clearly.",
        ],
    )
    add_case(
        doc,
        "Meeting Recall File Naming",
        "The final naming rule became YYYY-MM-DD – Meeting Name.m4a. This made files sort naturally "
        "and made them recognizable in NotebookLM and file pickers. Rename was deferred because direct "
        "rename was unreliable and memory-heavy copy experiments were unsafe for long recordings.",
    )
    add_prompt(
        doc,
        "Prompt: File Save Reliability",
        "Audit the save/export/share pipeline. Before marking a file saved, verify file exists, "
        "file size is greater than 0, filename extension is correct, and playback/readability can "
        "initialize if feasible. If validation fails, do not show success. Update QA docs with retest steps.",
    )

    add_h1(doc, "Part 9: Onboarding That Teaches Without Overexplaining", "The first minute should make the product feel obvious.")
    add_bullets(
        doc,
        [
            "Explain what the app does.",
            "Explain what the app does not do.",
            "Set expectations for external services.",
            "Explain local/privacy behavior if relevant.",
            "Prepare the user for permissions without pressuring them.",
            "Keep every screen short.",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Onboarding Flow",
        "Create a 3-screen onboarding flow for this app. Each screen should have one headline and one short body. "
        "Avoid technical jargon. Clarify the workflow, privacy expectations, and any external service relationship.",
    )

    add_h1(doc, "Part 10: Beta Testing Like A Founder", "Use beta testing to find trust breaks, not just bugs.")
    add_bullets(
        doc,
        [
            "Where does the tester hesitate?",
            "Do they understand the main workflow without explanation?",
            "Do they trust the result?",
            "Can they recover from mistakes?",
            "Do they know where files or outputs go?",
            "Do issues vary by device or platform?",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Beta Testing Guide",
        "Create a beta testing guide for this app. Include ideal tester profiles, what they should test, what I "
        "should observe, feedback questions, known limitations, and what feedback to prioritize or ignore.",
    )
    add_case(
        doc,
        "Closed Testing Answers",
        "When applying for Google Play production access, the strongest answers were plain and honest: testers "
        "came from friends, family, and professional contacts; they tested recording, save, playback, share, and "
        "NotebookLM upload; feedback led to fixes around file access, permission wording, and handoff reliability.",
    )

    add_h1(doc, "Part 11: Store Submission Is Product Work", "App review is not just paperwork.")
    add_checklist(
        doc,
        "Store Readiness Checklist",
        [
            "App icon is final enough for launch.",
            "Splash screen is clean.",
            "Screenshots match the actual app.",
            "Description does not overpromise.",
            "Privacy policy URL works.",
            "Support URL works.",
            "Data safety/privacy answers match reality.",
            "Permission purpose strings are clear and not manipulative.",
            "Test account or review notes are provided if needed.",
            "Core flow has been tested from a fresh install.",
            "No debug UI is visible.",
        ],
    )
    add_h2(doc, "App Store Review Lessons")
    add_bullets(
        doc,
        [
            "Apple may ask for a video showing the real app flow on a physical device.",
            "Apple can reject permission screens that use “Allow” before the system prompt.",
            "Review notes should explain external services clearly.",
            "If your app uses AI through another service, explain that your app does not generate the AI output itself.",
            "If login is optional, say so clearly.",
            "If data stays local, explain that accurately in privacy responses.",
        ],
    )
    add_h2(doc, "Google Play Review Lessons")
    add_bullets(
        doc,
        [
            "New personal developer accounts may need closed testing before production access.",
            "Google may ask how testers were recruited and what feedback changed.",
            "Be honest and specific. Do not inflate numbers.",
            "Production readiness answers should mention tested flows and fixed issues.",
            "Release notes should be short and user-facing.",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Store Rejection Response",
        "Apple/Google rejected my app for this reason: [paste message]. Explain the root cause in plain English, "
        "identify the app changes needed, update any relevant copy/docs, and draft a concise response to the reviewer.",
    )

    add_h1(doc, "Part 12: Working With Codex Day To Day", "The operating rhythm that keeps the project sane.")
    add_bullets(
        doc,
        [
            "Start each task with the exact user problem.",
            "Tell Codex what not to change.",
            "Ask it to inspect the code before editing.",
            "Ask for focused changes and verification.",
            "Commit after meaningful stable checkpoints.",
            "Keep docs updated with behavior changes.",
            "When a bug appears, show screenshots and exact observed behavior.",
            "Use real-device feedback as product evidence, not annoyance.",
        ],
    )
    add_prompt(
        doc,
        "Prompt: Stabilize After A Bad Spike",
        "The current spike has become unstable. Stop adding functionality. Restore the simplest working flow: "
        "[describe flow]. Remove or disable experimental code that risks reliability. Keep only the validated "
        "behavior. Update docs with what was rolled back and what is deferred.",
    )
    add_prompt(
        doc,
        "Prompt: Production Cleanup",
        "Perform a production cleanup pass. Remove debug UI, spike remnants, test controls, raw diagnostics, "
        "and unused experimental code. Preserve core behavior. Run checks. Create cleanup notes with what was "
        "removed, what remains dev-only, and known technical debt.",
    )

    add_h1(doc, "Part 13: Prompt Library", "Copy, paste, and adapt these prompts.")
    prompts = [
        ("Idea Audit", "Review this app idea: [idea]. Identify the target user, core problem, emotional reason to exist, MVP, risks, and what should be out of scope. Ask me any critical questions before recommending a build plan."),
        ("Product Scope Guard", "Create a PRODUCT_LOCK.md that keeps this app focused. Include what it is, what it is not, core workflow, MVP, out-of-scope features, and success definition."),
        ("UI Gap Tracker", "Audit the UI for missing states, edge cases, and trust gaps. Create a prioritized UI_GAP_TRACKER.md with P0/P1/P2 items across onboarding, home, core workflow, detail screens, settings, accessibility, and error states."),
        ("Screen Matrix", "Create a SCREEN_STATE_MATRIX.md. Define default, loading, empty, success, error, and interrupted states for every major screen."),
        ("Architecture Decisions", "Create ARCHITECTURE_DECISIONS.md. Document the technical philosophy, dependency rules, state management approach, platform constraints, and what complexity to avoid."),
        ("Stack Validation", "Review the current codebase and validate whether the stack can support the product requirements. Be direct about limitations, native module needs, app store readiness, and recommended path forward."),
        ("Implementation Phases", "Create MVP_IMPLEMENTATION_PHASES.md based on what has been technically validated. Define what ships in MVP, what is deferred, technical priorities, UX priorities, and implementation order."),
        ("Bug Fix", "Observed issue: [issue]. Reproduce or inspect the relevant code. Fix only the issue preventing the flow from working. Do not add new features. Run available checks and explain root cause."),
        ("App Store Rejection", "Read this app store rejection: [paste]. Identify the exact policy issue, required app changes, any copy changes, any config changes, and draft a reviewer response. Keep the fix minimal and compliant."),
        ("Release Notes", "Write a release name and concise release notes for this build. Mention user-facing fixes only. Keep the tone clear and professional."),
    ]
    for title, prompt in prompts:
        add_prompt(doc, title, prompt)

    add_h1(doc, "Part 14: Final Launch Checklist", "Use this before every real submission.")
    add_checklist(
        doc,
        "Product",
        [
            "The app has one clear core workflow.",
            "The main CTA is obvious.",
            "The app does not overpromise.",
            "The user always knows what to do next.",
            "Destructive actions require confirmation.",
            "Empty and error states are calm and useful.",
        ],
    )
    add_checklist(
        doc,
        "Technical",
        [
            "Typecheck/build passes.",
            "Core flow works on real devices.",
            "Permissions work from a fresh install.",
            "Files/data persist after restart.",
            "Share/export/upload flows work where required.",
            "No debug UI is visible.",
            "Known limitations are documented.",
        ],
    )
    add_checklist(
        doc,
        "Store",
        [
            "Screenshots are current and not misleading.",
            "App icon and splash are configured.",
            "Privacy policy and support URLs work.",
            "Data safety/privacy answers are accurate.",
            "Permission purpose strings are neutral and clear.",
            "Review notes explain how to test.",
            "Version/build numbers are correct.",
            "Release notes are ready.",
        ],
    )

    add_h1(doc, "Final Word")
    add_p(
        doc,
        "The best builders are not the people who ask AI for the most code. They are "
        "the people who ask better questions, protect the user experience, test the "
        "riskiest assumptions early, and keep the product focused when it would be easier "
        "to add more.",
    )
    add_p(
        doc,
        "Codex and ChatGPT can help you move incredibly fast. But speed only matters if "
        "you are moving in the right direction. Use AI to clarify, build, test, revise, "
        "document, and ship. Do not use it to avoid thinking.",
    )


def build():
    doc = Document()
    setup_styles(doc)
    add_front_matter(doc)
    add_main_content(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
