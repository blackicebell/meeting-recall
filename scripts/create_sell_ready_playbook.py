from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUTPUT = Path(
    r"C:\Users\creat\Documents\app meeting recall\docs"
    r"\The AI App Building Playbook - Sell Ready Edition.docx"
)


def shade(cell, fill):
    props = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    props.append(shd)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(7)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size in [("Heading 1", 22), ("Heading 2", 15), ("Heading 3", 12)]:
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].paragraph_format.space_before = Pt(12)
        styles[name].paragraph_format.space_after = Pt(6)

    styles["Subtitle"].font.name = "Aptos"
    styles["Subtitle"].font.size = Pt(12)
    styles["Subtitle"].font.italic = True


def p(doc, text=""):
    return doc.add_paragraph(text)


def h1(doc, text, subtitle=None):
    doc.add_page_break()
    doc.add_paragraph(text, style="Heading 1")
    if subtitle:
        doc.add_paragraph(subtitle, style="Subtitle")


def h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def h3(doc, text):
    doc.add_paragraph(text, style="Heading 3")


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def prompt(doc, title, text):
    h3(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "F3F6FC")
    run = cell.paragraphs[0].add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    doc.add_paragraph()


def lesson(doc, title, text):
    h3(doc, title)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, "FFF7F7")
    cell.paragraphs[0].add_run(text)
    doc.add_paragraph()


def checklist(doc, title, items):
    h3(doc, title)
    for item in items:
        line = doc.add_paragraph(style="List Bullet")
        line.add_run("[ ] ").bold = True
        line.add_run(item)


def title_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("The AI App Building Playbook")
    run.bold = True
    run.font.size = Pt(30)
    run.font.name = "Aptos Display"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "A step-by-step guide for turning any app idea into a real product with ChatGPT, Codex, and AI-assisted development"
    )
    run.font.size = Pt(13)

    edition = doc.add_paragraph()
    edition.alignment = WD_ALIGN_PARAGRAPH.CENTER
    edition.add_run("From first idea to app store launch").italic = True

    doc.add_page_break()
    h2(doc, "Who This Book Is For")
    p(
        doc,
        "This book is for founders, designers, creators, students, solo builders, and beginners who have an app idea but do not know how to turn it into something real.",
    )
    p(
        doc,
        "You do not need to be a developer to use this process. You do need to be willing to think clearly, ask better questions, test the product, and move one focused step at a time.",
    )
    p(
        doc,
        "The goal is simple: help you use ChatGPT and Codex as practical product-building partners, from the earliest idea all the way to Google Play and the Apple App Store.",
    )

    h2(doc, "How To Use This Book")
    p(
        doc,
        "Do not treat this as theory. Treat it like a working playbook. Every section should help you produce something useful: a clearer idea, a better prompt, a product brief, a screen map, a risk list, a QA plan, or a launch checklist.",
    )
    numbers(
        doc,
        [
            "Start with your idea in one plain sentence.",
            "Use the prompts to clarify the product before building.",
            "Create project documents that guide Codex before asking it to write code.",
            "Build in small phases instead of giant feature dumps.",
            "Validate risky technical workflows before polishing the UI.",
            "Test on real devices before trusting the app.",
            "Treat app store review as part of product design, not paperwork at the end.",
        ],
    )


def build_content(doc):
    h1(doc, "Part 1: Start With The Product, Not The Code", "AI works best when you give it a clear target.")
    h2(doc, "The Biggest Mistake New Builders Make")
    p(
        doc,
        "Most people ask AI to build an app too early. That usually creates a messy prototype because the AI has no real boundaries. It does not know who the app is for, what matters most, what should be avoided, or which technical risks could sink the project.",
    )
    p(
        doc,
        "The better approach is to use AI as a product partner before you use it as a coding tool. Make it help you think. Make it ask questions. Make it challenge the scope. Make it identify risky assumptions.",
    )
    h2(doc, "The App Clarity Stack")
    bullets(
        doc,
        [
            "Problem: What pain does this app reduce?",
            "Audience: Who feels that pain often enough to care?",
            "Moment: When does the user reach for this app?",
            "Outcome: What does the user get after using it?",
            "Trust: What must never go wrong?",
            "MVP: What is the smallest version that still delivers the core outcome?",
            "Not Included: What are you intentionally refusing to build?",
        ],
    )
    prompt(
        doc,
        "Prompt: Clarify Your Idea",
        "I have an app idea: [describe it]. Help me clarify the problem, target user, emotional reason to exist, core workflow, MVP, and what I should avoid building too early. Be direct and challenge weak assumptions.",
    )
    prompt(
        doc,
        "Prompt: Define What Must Never Break",
        "For this app idea, identify the 5 things that must never go wrong for users to trust it. For each one, explain the product risk, technical risk, and UX risk.",
    )
    lesson(
        doc,
        "Field Lesson: A Strong App Knows What It Is Not",
        "A focused app becomes easier to build when you define what it refuses to become. Before adding features, write down the products your app is not trying to be. This protects the core workflow and helps Codex avoid building in the wrong direction.",
    )

    h1(doc, "Part 2: Create The Documents That Guide Codex", "A prompt is temporary. Project rules are durable.")
    h2(doc, "Why Documentation Matters")
    p(
        doc,
        "AI coding tools perform better when they have durable context. If you care about scope, UX, technical decisions, launch rules, or review compliance, write them down inside the project.",
    )
    h2(doc, "Core Documents To Create")
    bullets(
        doc,
        [
            "PRODUCT_LOCK.md: what the product is and is not.",
            "UX_RULES.md: copy, flow, emotional tone, and friction rules.",
            "COMPONENT_SYSTEM.md: buttons, typography, spacing, states, and icons.",
            "TECHNICAL_EXPECTATIONS.md: platform behavior, storage, permissions, APIs, and reliability rules.",
            "IMPLEMENTATION_PLAN.md: phases and build order.",
            "QA_TEST_PLAN.md: real flows to test before launch.",
            "FEATURE_DONE_CHECKLIST.md: what done actually means.",
            "SCREEN_STATE_MATRIX.md: loading, empty, success, error, and interrupted states.",
            "ARCHITECTURE_DECISIONS.md: why you chose one technical path over another.",
        ],
    )
    prompt(
        doc,
        "Prompt: Create A Product Lock",
        "Create PRODUCT_LOCK.md for this app. Define the app, target user, core workflow, MVP scope, out-of-scope features, success definition, and product principles. Be strict about avoiding feature creep.",
    )
    prompt(
        doc,
        "Prompt: Create A QA Plan",
        "Create QA_TEST_PLAN.md for this app. Include first launch, permissions, core workflow, saved item flow, old item flow, share/export flow, error states, app restart, and real-device testing. Make it practical for beta testers.",
    )

    h1(doc, "Part 3: Prompt Codex Like A Product Lead", "Good prompts reduce chaos.")
    h2(doc, "The Codex Prompt Formula")
    bullets(
        doc,
        [
            "Goal: what should be true when the task is done.",
            "Do not: what Codex must avoid changing or adding.",
            "Requirements: specific behavior, copy, screens, and edge cases.",
            "References: docs or files Codex should follow.",
            "Validation: tests, typecheck, build, or manual QA steps.",
            "Summary: what Codex should report back.",
        ],
    )
    prompt(
        doc,
        "Reusable Codex Task Prompt",
        "Goal:\n[Describe the outcome.]\n\nDo not:\n- Do not add unrelated features.\n- Do not refactor unrelated code.\n- Do not change working flows unless needed.\n\nRequirements:\n1. [Requirement]\n2. [Requirement]\n3. [Requirement]\n\nFollow these docs:\n- /docs/PRODUCT_LOCK.md\n- /docs/UX_RULES.md\n- /docs/TECHNICAL_EXPECTATIONS.md\n\nAfter completion:\n- run typecheck/build if available\n- summarize files changed\n- explain what to test",
    )
    h2(doc, "Rules That Save Projects")
    bullets(
        doc,
        [
            "Ask for one meaningful change at a time.",
            "Use “Do not implement...” to prevent accidental scope creep.",
            "Tell Codex what must be preserved.",
            "Ask it to update docs when behavior changes.",
            "Ask for exact testing steps after every major change.",
            "When something breaks, ask for root cause before adding features.",
            "When a flow feels wrong, describe the emotional problem, not just the UI bug.",
        ],
    )

    h1(doc, "Part 4: Build In Phases, Not Feature Piles", "A clean build order prevents expensive mistakes.")
    h2(doc, "Recommended Build Order")
    numbers(
        doc,
        [
            "Project audit and stack validation.",
            "App shell and navigation.",
            "Theme and reusable components.",
            "Core screens with placeholders.",
            "Technical spikes for risky workflows.",
            "Production version of the core workflow.",
            "Saved item/history flow.",
            "Primary handoff/share/export flow.",
            "Settings, onboarding, and permissions.",
            "Error states and recovery states.",
            "Beta testing cleanup.",
            "Store assets, screenshots, and submission.",
        ],
    )
    checklist(
        doc,
        "Phase Gate Checklist",
        [
            "Can the user complete the core workflow end to end?",
            "Did we test on a real device?",
            "Did we document known limitations?",
            "Did we update QA steps?",
            "Are debug screens or spike controls still visible?",
            "Does the app still match the product lock?",
            "Is this change worth the added complexity?",
        ],
    )

    h1(doc, "Part 5: Validate Risk Before Polishing UI", "Technical spikes protect your launch.")
    p(
        doc,
        "A spike is a small technical experiment that answers a risky question. It is not final UI. It is how you avoid spending weeks designing around an assumption that turns out to be false.",
    )
    bullets(
        doc,
        [
            "Can the app record audio reliably on real devices?",
            "Can files be saved somewhere users can find?",
            "Can files be shared with the correct MIME type?",
            "Can the target external app accept the file?",
            "Can permissions be requested without creating review problems?",
            "Does the platform support the flow you imagined?",
        ],
    )
    lesson(
        doc,
        "Field Lesson: Spikes Should Be Disposable",
        "A spike is allowed to look ugly. Its job is to answer a question. Once the risk is understood, either turn the working part into production code or throw the spike away. Do not let debug buttons and experimental paths become the real product.",
    )
    prompt(
        doc,
        "Prompt: Define Technical Spikes",
        "Identify the highest-risk technical assumptions in this app. Create IMPLEMENTATION_SPIKES.md with each spike goal, test steps, pass/fail conditions, platform risks, and what decision the spike should unlock. Order the spikes by product risk.",
    )

    h1(doc, "Part 6: UX Is Trust, Not Decoration", "Design the moments where users might doubt the app.")
    p(
        doc,
        "A user-friendly app does not just look clean. It reduces doubt. The user should not wonder what to do next, whether something saved, where a file went, or whether a button is safe to press.",
    )
    bullets(
        doc,
        [
            "Loading: what is happening right now?",
            "Empty: what can I do next?",
            "Success: did it work?",
            "Error: what went wrong in plain language?",
            "Interrupted: did I lose anything?",
            "Recovery: what action can I take now?",
        ],
    )
    lesson(
        doc,
        "Field Lesson: Use Readiness Gates For Critical Actions",
        "If an action depends on a file, payment, account, upload, permission, or external service, do not activate the button too early. Show a calm preparing state, verify the requirement, then enable the action when the app is truly ready.",
    )
    lesson(
        doc,
        "Field Lesson: Permission Copy Can Affect Approval",
        "On iOS, custom pre-permission screens should provide context without pressuring the user. A neutral button like “Continue” is safer than a button that says “Allow” before the native system prompt appears.",
    )
    prompt(
        doc,
        "Prompt: Add A Readiness Gate",
        "Add a readiness gate before enabling [critical action]. The screen should show “Preparing...” while validation runs, enable the primary action only after validation passes, and show a calm failed state with Try Again if validation fails. Do not add normal-state warning copy.",
    )

    h1(doc, "Part 7: Understand Native Mobile Reality", "Real devices reveal what simulators hide.")
    p(
        doc,
        "Simulators are useful, but they do not fully represent file pickers, permissions, audio interruptions, storage providers, background behavior, or app review conditions. If the app depends on microphone, storage, camera, Bluetooth, notifications, calendar, or location, test on real devices early.",
    )
    h2(doc, "Android vs iOS Differences")
    bullets(
        doc,
        [
            "Android has many device brands, file providers, storage behaviors, and picker inconsistencies.",
            "iOS is more controlled, but sandboxing can make files less browsable unless sharing/export is designed well.",
            "Android may show a file but another app may gray it out due to MIME or provider behavior.",
            "iOS often prefers share/open-in flows over expecting users to browse app storage.",
            "Both platforms need clear permission copy and recovery states.",
        ],
    )
    prompt(
        doc,
        "Prompt: Stack Validation",
        "Review the current codebase and validate whether the stack can support the product requirements. Be direct about limitations, native module needs, app store readiness, and recommended path forward.",
    )

    h1(doc, "Part 8: Files, Sharing, And External Workflows", "If files are part of the product, files are part of the UX.")
    p(
        doc,
        "For apps that record, export, import, or upload files, storage is part of the user experience. Users need confidence that the file exists, has the right name, can be found, and can be used by the target app.",
    )
    bullets(
        doc,
        [
            "Use human-readable filenames.",
            "Create the correct filename during initial save whenever possible.",
            "Validate file size before marking a save successful.",
            "Validate playback or readability when feasible.",
            "Do not create fake success states.",
            "Make Share a reliable fallback for platform picker weirdness.",
            "Document platform differences clearly.",
        ],
    )
    prompt(
        doc,
        "Prompt: File Save Reliability",
        "Audit the save/export/share pipeline. Before marking a file saved, verify file exists, file size is greater than 0, filename extension is correct, and playback/readability can initialize if feasible. If validation fails, do not show success. Update QA docs with retest steps.",
    )

    h1(doc, "Part 9: Onboarding That Teaches Without Overexplaining", "The first minute should make the product feel obvious.")
    bullets(
        doc,
        [
            "Explain what the app does.",
            "Explain what the app does not do.",
            "Set expectations for external services.",
            "Explain local/privacy behavior if relevant.",
            "Prepare the user for permissions without pressuring them.",
            "Keep every screen short.",
        ],
    )
    prompt(
        doc,
        "Prompt: Onboarding Flow",
        "Create a 3-screen onboarding flow for this app. Each screen should have one headline and one short body. Avoid technical jargon. Clarify the workflow, privacy expectations, and any external service relationship.",
    )

    h1(doc, "Part 10: Beta Testing Like A Founder", "Use beta testing to find trust breaks, not just bugs.")
    bullets(
        doc,
        [
            "Where does the tester hesitate?",
            "Do they understand the main workflow without explanation?",
            "Do they trust the result?",
            "Can they recover from mistakes?",
            "Do they know where files or outputs go?",
            "Do issues vary by device or platform?",
        ],
    )
    prompt(
        doc,
        "Prompt: Beta Testing Guide",
        "Create a beta testing guide for this app. Include ideal tester profiles, what they should test, what I should observe, feedback questions, known limitations, and what feedback to prioritize or ignore.",
    )

    h1(doc, "Part 11: Store Submission Is Product Work", "App review is part of the product.")
    checklist(
        doc,
        "Store Readiness Checklist",
        [
            "App icon is final enough for launch.",
            "Splash screen is clean.",
            "Screenshots match the actual app.",
            "Description does not overpromise.",
            "Privacy policy URL works.",
            "Support URL works.",
            "Data safety/privacy answers match reality.",
            "Permission purpose strings are clear and not manipulative.",
            "Test account or review notes are provided if needed.",
            "Core flow has been tested from a fresh install.",
            "No debug UI is visible.",
        ],
    )
    h2(doc, "App Store Review Lessons")
    bullets(
        doc,
        [
            "Apple may ask for a video showing the real app flow on a physical device.",
            "Apple can reject permission screens that use “Allow” before the system prompt.",
            "Review notes should explain external services clearly.",
            "If your app uses AI through another service, explain that your app does not generate the AI output itself.",
            "If login is optional, say so clearly.",
            "If data stays local, explain that accurately in privacy responses.",
        ],
    )
    h2(doc, "Google Play Review Lessons")
    bullets(
        doc,
        [
            "New personal developer accounts may need closed testing before production access.",
            "Google may ask how testers were recruited and what feedback changed.",
            "Be honest and specific. Do not inflate numbers.",
            "Production readiness answers should mention tested flows and fixed issues.",
            "Release notes should be short and user-facing.",
        ],
    )
    prompt(
        doc,
        "Prompt: Store Rejection Response",
        "Apple/Google rejected my app for this reason: [paste message]. Explain the root cause in plain English, identify the app changes needed, update any relevant copy/docs, and draft a concise response to the reviewer.",
    )

    h1(doc, "Part 12: Working With Codex Day To Day", "The operating rhythm that keeps the project sane.")
    bullets(
        doc,
        [
            "Start each task with the exact user problem.",
            "Tell Codex what not to change.",
            "Ask it to inspect the code before editing.",
            "Ask for focused changes and verification.",
            "Commit after meaningful stable checkpoints.",
            "Keep docs updated with behavior changes.",
            "When a bug appears, show screenshots and exact observed behavior.",
            "Use real-device feedback as product evidence, not annoyance.",
        ],
    )
    prompt(
        doc,
        "Prompt: Stabilize After A Bad Spike",
        "The current spike has become unstable. Stop adding functionality. Restore the simplest working flow: [describe flow]. Remove or disable experimental code that risks reliability. Keep only the validated behavior. Update docs with what was rolled back and what is deferred.",
    )
    prompt(
        doc,
        "Prompt: Production Cleanup",
        "Perform a production cleanup pass. Remove debug UI, spike remnants, test controls, raw diagnostics, and unused experimental code. Preserve core behavior. Run checks. Create cleanup notes with what was removed, what remains dev-only, and known technical debt.",
    )

    h1(doc, "Part 13: Prompt Library", "Copy, paste, and adapt these prompts.")
    for title, body in [
        ("Idea Audit", "Review this app idea: [idea]. Identify the target user, core problem, emotional reason to exist, MVP, risks, and what should be out of scope. Ask me any critical questions before recommending a build plan."),
        ("Product Scope Guard", "Create a PRODUCT_LOCK.md that keeps this app focused. Include what it is, what it is not, core workflow, MVP, out-of-scope features, and success definition."),
        ("UI Gap Tracker", "Audit the UI for missing states, edge cases, and trust gaps. Create a prioritized UI_GAP_TRACKER.md with P0/P1/P2 items across onboarding, home, core workflow, detail screens, settings, accessibility, and error states."),
        ("Screen Matrix", "Create a SCREEN_STATE_MATRIX.md. Define default, loading, empty, success, error, and interrupted states for every major screen."),
        ("Architecture Decisions", "Create ARCHITECTURE_DECISIONS.md. Document the technical philosophy, dependency rules, state management approach, platform constraints, and what complexity to avoid."),
        ("Stack Validation", "Review the current codebase and validate whether the stack can support the product requirements. Be direct about limitations, native module needs, app store readiness, and recommended path forward."),
        ("Implementation Phases", "Create MVP_IMPLEMENTATION_PHASES.md based on what has been technically validated. Define what ships in MVP, what is deferred, technical priorities, UX priorities, and implementation order."),
        ("Bug Fix", "Observed issue: [issue]. Reproduce or inspect the relevant code. Fix only the issue preventing the flow from working. Do not add new features. Run available checks and explain root cause."),
        ("App Store Rejection", "Read this app store rejection: [paste]. Identify the exact policy issue, required app changes, any copy changes, any config changes, and draft a concise response to the reviewer."),
        ("Release Notes", "Write a release name and concise release notes for this build. Mention user-facing fixes only. Keep the tone clear and professional."),
    ]:
        prompt(doc, title, body)

    h1(doc, "Part 14: Final Launch Checklist", "Use this before every real submission.")
    checklist(
        doc,
        "Product",
        [
            "The app has one clear core workflow.",
            "The main CTA is obvious.",
            "The app does not overpromise.",
            "The user always knows what to do next.",
            "Destructive actions require confirmation.",
            "Empty and error states are calm and useful.",
        ],
    )
    checklist(
        doc,
        "Technical",
        [
            "Typecheck/build passes.",
            "Core flow works on real devices.",
            "Permissions work from a fresh install.",
            "Files/data persist after restart.",
            "Share/export/upload flows work where required.",
            "No debug UI is visible.",
            "Known limitations are documented.",
        ],
    )
    checklist(
        doc,
        "Store",
        [
            "Screenshots are current and not misleading.",
            "App icon and splash are configured.",
            "Privacy policy and support URLs work.",
            "Data safety/privacy answers are accurate.",
            "Permission purpose strings are neutral and clear.",
            "Review notes explain how to test.",
            "Version/build numbers are correct.",
            "Release notes are ready.",
        ],
    )

    h1(doc, "Final Word")
    p(
        doc,
        "The best builders are not the people who ask AI for the most code. They are the people who ask better questions, protect the user experience, test the riskiest assumptions early, and keep the product focused when it would be easier to add more.",
    )
    p(
        doc,
        "Codex and ChatGPT can help you move incredibly fast. But speed only matters if you are moving in the right direction. Use AI to clarify, build, test, revise, document, and ship. Do not use it to avoid thinking.",
    )


def build():
    doc = Document()
    setup(doc)
    title_page(doc)
    build_content(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
