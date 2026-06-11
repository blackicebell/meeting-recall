from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Android_App_Subscription_Monetization_Guide.docx"

BLUE = RGBColor(75, 125, 230)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(100, 116, 139)
LIGHT_BLUE = "EAF1FF"
LIGHT_GRAY = "F8FAFC"
BORDER = "D7DEE8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_paragraph(paragraph, before=0, after=6, line_spacing=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def add_title(doc):
    p = doc.add_paragraph()
    style_paragraph(p, after=4)
    run = p.add_run("Android App Subscription Monetization Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(25)
    run.font.bold = True
    run.font.color.rgb = INK

    p = doc.add_paragraph()
    style_paragraph(p, after=12)
    run = p.add_run(
        "A practical checklist for setting up Google Play subscriptions with RevenueCat."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    add_callout(
        doc,
        "Recommended model",
        "Use a free app download with an in-app subscription. Avoid switching an already-free Google Play app into a paid-upfront app, because that usually requires creating a new app listing.",
    )


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    style_paragraph(p, before=14 if level == 1 else 10, after=6)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BLUE if level <= 2 else INK
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = INK
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, after=3)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style_paragraph(p, after=3)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = INK


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_border(cell, "C7D7FF")
    p = cell.paragraphs[0]
    style_paragraph(p, after=3)
    r = p.add_run(title)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = INK
    p = cell.add_paragraph()
    style_paragraph(p, after=0)
    r = p.add_run(body)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    doc.add_paragraph()


def add_two_col_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_width(table, [1.8, 4.7])
    hdr = table.rows[0].cells
    hdr[0].text = "Area"
    hdr[1].text = "What to check"
    for cell in hdr:
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_border(cell)
        for p in cell.paragraphs:
            style_paragraph(p, after=0)
            for r in p.runs:
                r.font.name = "Calibri"
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = INK

    for label, detail in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = detail
        for cell in cells:
            set_cell_border(cell)
            for p in cell.paragraphs:
                style_paragraph(p, after=0)
                for r in p.runs:
                    r.font.name = "Calibri"
                    r.font.size = Pt(10)
                    r.font.color.rgb = INK
    doc.add_paragraph()


def add_code_line(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=2, after=6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = INK


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK

    add_title(doc)

    add_heading(doc, "1. Choose the monetization model")
    add_body(
        doc,
        "For most modern mobile apps, the safest model is a free download with paid Pro features unlocked by subscription. This keeps acquisition friction low while still giving serious users a path to pay.",
    )
    add_bullets(
        doc,
        [
            "Free download: users can install and try the app first.",
            "Subscription: unlocks premium or unlimited behavior inside the app.",
            "Avoid paid-upfront unless the app's value is obvious before install.",
        ]
    )

    add_heading(doc, "2. Add billing support to the app")
    add_body(
        doc,
        "Google Play may not unlock subscription setup until it sees a build that includes billing support. With RevenueCat, the billing layer is handled through the RevenueCat SDK.",
    )
    add_code_line(doc, "npx expo install react-native-purchases react-native-purchases-ui")
    add_bullets(
        doc,
        [
            "Use the Android RevenueCat public SDK key, usually starting with goog_.",
            "Do not place Google service account JSON inside app code.",
            "Do not use the iOS RevenueCat key for Android builds.",
        ]
    )

    add_heading(doc, "3. Create and upload a production Android build")
    add_body(
        doc,
        "After the billing SDK is installed and configured, create a new Android App Bundle and upload it to Google Play. Internal testing is usually the safest first destination.",
    )
    add_code_line(doc, "eas build -p android --profile production")
    add_bullets(
        doc,
        [
            "Upload the .aab to Internal testing, Closed testing, or Production.",
            "Wait for Google Play to process the build.",
            "Return to Products > Subscriptions after processing is complete.",
        ]
    )

    add_heading(doc, "4. Create subscriptions in Google Play Console")
    add_body(
        doc,
        "In Google Play Console, create stable subscription product IDs. Keep product IDs simple because changing them later can create avoidable cleanup work.",
    )
    add_two_col_table(
        doc,
        [
            ("Monthly", "Product ID: monthly. Billing period: 1 month. Example price: $3.99/month. Optional free trial: 7 days."),
            ("Yearly", "Product ID: yearly. Billing period: 1 year. Example price: $29.99/year. Optional free trial: 7 days."),
        ],
    )

    add_heading(doc, "5. Create the Android app in RevenueCat")
    add_body(
        doc,
        "In RevenueCat, add a Google Play app under the same project. The package name must match the Android package name exactly.",
    )
    add_bullets(
        doc,
        [
            "App type: Google Play.",
            "Package name: com.yourapp.package.",
            "Copy the public SDK key that starts with goog_.",
            "Use that key only for Android.",
        ]
    )

    add_heading(doc, "6. Create Google service account credentials")
    add_body(
        doc,
        "RevenueCat needs a Google service account JSON file to validate purchases and subscription status with Google Play.",
    )
    add_numbered(
        doc,
        [
            "Create a service account in Google Cloud.",
            "Enable Google Play Android Developer API, Google Play Developer Reporting API, and Cloud Pub/Sub API.",
            "Download a JSON key for the service account.",
            "Invite the service account email in Google Play Console under Users and permissions.",
            "Grant app-level permissions for app information, financial data, orders/subscriptions, and store presence.",
            "Upload the JSON file to RevenueCat and save.",
        ]
    )
    add_callout(
        doc,
        "Do not ship secrets",
        "The service account JSON belongs in RevenueCat only. Never include it in your app code, repository, or public documentation.",
    )

    add_heading(doc, "7. Import products into RevenueCat")
    add_body(
        doc,
        "Once Google Play products exist and credentials are valid, import or create matching products in RevenueCat.",
    )
    add_bullets(
        doc,
        [
            "Import monthly and yearly under the Google Play app section.",
            "Attach both products to the same Pro entitlement.",
            "Keep entitlement naming stable and consistent across platforms.",
        ]
    )

    add_heading(doc, "8. Create the default offering")
    add_body(
        doc,
        "Offerings decide what the paywall presents. Most apps can start with one default offering that contains monthly and annual packages.",
    )
    add_two_col_table(
        doc,
        [
            ("Offering ID", "default"),
            ("Monthly package", "Google Play product: monthly"),
            ("Annual package", "Google Play product: yearly"),
            ("Entitlement", "Pro or your app-specific Pro entitlement"),
        ],
    )

    add_heading(doc, "9. Build the paywall")
    add_body(
        doc,
        "The paywall should use store-provided prices instead of hardcoded text. This prevents pricing mismatches across countries and review environments.",
    )
    add_bullets(
        doc,
        [
            "Show subscription title, duration, price, and trial clearly.",
            "Include Restore Purchases.",
            "Include Terms and Privacy links where required.",
            "Do not exaggerate what the subscription provides.",
        ]
    )

    add_heading(doc, "10. Test with Google Play testers")
    add_body(
        doc,
        "Always test from a Google Play-installed build, not just a local dev build. Subscription behavior depends on the Play Store environment.",
    )
    add_numbered(
        doc,
        [
            "Add tester accounts in Google Play Console.",
            "Install the app from the test track link.",
            "Open the paywall and confirm products load.",
            "Verify monthly and yearly prices.",
            "Complete a sandbox purchase.",
            "Confirm Pro unlocks immediately.",
            "Close and reopen the app to confirm access persists.",
            "Test Restore Purchases.",
        ]
    )

    add_heading(doc, "11. Avoid common mistakes")
    add_two_col_table(
        doc,
        [
            ("Wrong SDK key", "Android must use the RevenueCat key that starts with goog_. iOS uses the key that starts with appl_."),
            ("Wrong products", "Do not leave the paywall connected to RevenueCat Test Store products for production."),
            ("Hardcoded prices", "Use RevenueCat/store pricing so App Review and users see accurate localized prices."),
            ("Missing build", "Google Play may require a processed build with billing support before subscriptions can be created."),
            ("Bad user trust", "Do not block users from accessing content they already created."),
        ],
    )

    add_heading(doc, "12. Production checklist")
    add_bullets(
        doc,
        [
            "Android build includes billing SDK.",
            "Google Play subscriptions are created and active for testing.",
            "RevenueCat Android app has valid service account credentials.",
            "Products are imported into RevenueCat.",
            "Products are attached to the correct entitlement.",
            "Default offering contains monthly and yearly packages.",
            "App uses platform-specific RevenueCat SDK keys.",
            "Paywall shows correct prices.",
            "Purchase unlocks Pro.",
            "Restore Purchases works.",
            "Existing user data remains accessible.",
        ]
    )

    add_callout(
        doc,
        "Simple launch stance",
        "Keep the free app useful, make Pro clearly worth paying for, and avoid taking away access to anything users already created.",
    )

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
